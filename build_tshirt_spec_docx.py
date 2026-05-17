from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

PURPLE = RGBColor(0x7B, 0x2F, 0xBE)
GOLD = RGBColor(0xF5, 0xC8, 0x42)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY = RGBColor(0x1A, 0x1A, 0x1A)
LIGHT_PURPLE = RGBColor(0x6A, 0x1F, 0xAA)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_header(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, '7B2FBE')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  {text}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    doc.add_paragraph()

def add_data_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, '1A1A1A')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = GOLD
    # Data rows
    for ri, row in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'F5F5F5'
        for ci, val in enumerate(row):
            cell = table.cell(ri+1, ci)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GREY
    # Set column widths
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])
    doc.add_paragraph()

# ── HEADER ──────────────────────────────────────────────────
header_table = doc.add_table(rows=3, cols=1)
header_table.style = 'Table Grid'

# Row 1 - Brand name
cell = header_table.cell(0, 0)
set_cell_bg(cell, '000000')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p.add_run('SACHi')
r1.font.bold = True
r1.font.size = Pt(28)
r1.font.color.rgb = PURPLE
r2 = p.add_run(' STREAM')
r2.font.bold = True
r2.font.size = Pt(28)
r2.font.color.rgb = GOLD

# Row 2 - Title
cell2 = header_table.cell(1, 0)
set_cell_bg(cell2, '000000')
p2 = cell2.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('T-SHIRT PRINT SPECIFICATION SHEET')
run2.font.bold = True
run2.font.size = Pt(13)
run2.font.color.rgb = WHITE

# Row 3 - Contact info
cell3 = header_table.cell(2, 0)
set_cell_bg(cell3, '000000')
p3 = cell3.paragraphs[0]
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('LDNA Consulting LLC  |  jaygnz27@gmail.com  |  908-255-2195  |  sachistream.com  |  May 17, 2026')
run3.font.size = Pt(8)
run3.font.color.rgb = GOLD

doc.add_paragraph()

# ── SECTION 1: GARMENT ──────────────────────────────────────
add_section_header(doc, '1.  GARMENT SPECIFICATIONS')
add_data_table(doc,
    ['FIELD', 'DETAIL'],
    [
        ['Garment Type', 'Unisex Crew Neck T-Shirt'],
        ['Fit', 'Fitted / Athletic cut'],
        ['Garment Color', 'Black — Pantone Black 6 C'],
        ['Fabric', '100% Cotton  OR  95% Cotton / 5% Elastane (athletic fit)'],
        ['Weight', '180–200 GSM recommended'],
        ['Sizes Required', 'S, M, L, XL, XXL'],
    ],
    [5, 12]
)

# ── SECTION 2: PRINT METHOD ─────────────────────────────────
add_section_header(doc, '2.  PRINT METHOD')
add_data_table(doc,
    ['METHOD', 'WHEN TO USE'],
    [
        ['DTF — Direct to Film  ★ RECOMMENDED', 'Best for gold detail, gradients, full color Dharma Chakra'],
        ['Screen Print', 'Use for large bulk quantities (50+ shirts)'],
    ],
    [7, 10]
)

# ── SECTION 3: FRONT PRINT ──────────────────────────────────
add_section_header(doc, '3.  FRONT PRINT DESIGN')
add_data_table(doc,
    ['FIELD', 'DETAIL'],
    [
        ['Position', 'Center chest — approx 3 inches (7.5 cm) below collar'],
        ['Print Size', '28 cm wide × 14 cm tall  (approx 11" × 5.5")'],
        ['"SACH" Text', 'Bold sans-serif, ALL CAPS — Purple #7B2FBE / Pantone 266 C'],
        ['"i" Letter', 'Lowercase, same font — Gold #F5C842 / Pantone 116 C'],
        ['"STREAM" Text', 'Smaller bold ALL CAPS, wide letter spacing — Gold #F5C842'],
        ['STREAM Size', 'Approx 40% height of "SACHi" — center aligned below'],
    ],
    [5, 12]
)

# ── SECTION 4: BACK PRINT ───────────────────────────────────
add_section_header(doc, '4.  BACK PRINT DESIGN')

p = doc.add_paragraph()
run = p.add_run('Element 1 — Dharma Chakra Logo Badge (upper back)')
run.font.bold = True
run.font.size = Pt(9)
run.font.color.rgb = PURPLE

add_data_table(doc,
    ['FIELD', 'DETAIL'],
    [
        ['Position', 'Upper center back — 3 inches (7.5 cm) below collar'],
        ['Print Size', '15 cm × 15 cm  (approx 6" × 6")'],
        ['Design', 'Ornate Buddhist Dharma Chakra — 8-spoke wheel, intricate gold carvings'],
        ['Wheel Color', 'Antique gold on black background'],
        ['Center Hub', 'Black circle with glowing gold ring border'],
        ['Hub Text Line 1', '"SACHi" — SACH in Purple #7B2FBE, "i" in Gold #F5C842'],
        ['Hub Text Line 2', '"STREAM" in small gold letters, center aligned'],
    ],
    [5, 12]
)

p = doc.add_paragraph()
run = p.add_run('Element 2 — Slogan')
run.font.bold = True
run.font.size = Pt(9)
run.font.color.rgb = PURPLE

add_data_table(doc,
    ['FIELD', 'DETAIL'],
    [
        ['Position', 'Center back — 2 cm below Dharma Chakra wheel'],
        ['Text', 'REAL PEOPLE. REAL VIDEOS. NO AI.'],
        ['Font', 'Bold ALL CAPS, clean sans-serif'],
        ['Color', 'Gold — Pantone 116 C / HEX #F5C842'],
        ['Letter Height', 'Approx 1.5 cm tall'],
    ],
    [5, 12]
)

p = doc.add_paragraph()
run = p.add_run('Element 3 — Website')
run.font.bold = True
run.font.size = Pt(9)
run.font.color.rgb = PURPLE

add_data_table(doc,
    ['FIELD', 'DETAIL'],
    [
        ['Position', 'Center back — 1.5 cm below slogan'],
        ['Text', 'sachistream.com'],
        ['Font', 'Regular weight sans-serif, lowercase'],
        ['Color', 'White — HEX #FFFFFF'],
        ['Letter Height', 'Approx 0.8 cm tall'],
    ],
    [5, 12]
)

# ── SECTION 5: COLOR REFERENCE ──────────────────────────────
add_section_header(doc, '5.  COLOR REFERENCE')
add_data_table(doc,
    ['COLOR NAME', 'PANTONE', 'HEX', 'CMYK', 'RGB'],
    [
        ['Purple', '266 C', '#7B2FBE', '73, 79, 0, 0', '123, 47, 190'],
        ['Gold', '116 C', '#F5C842', '0, 13, 80, 4', '245, 200, 66'],
        ['White', '—', '#FFFFFF', '0, 0, 0, 0', '255, 255, 255'],
        ['Black (garment)', 'Black 6 C', '#000000', '—', '0, 0, 0'],
    ],
    [3.5, 3, 3, 4, 3.5]
)

# ── SECTION 6: QUANTITIES ───────────────────────────────────
add_section_header(doc, '6.  QUANTITIES (TO BE CONFIRMED)')
add_data_table(doc,
    ['SIZE', 'QUANTITY'],
    [['S', ''], ['M', ''], ['L', ''], ['XL', ''], ['XXL', ''], ['TOTAL', '']],
    [4, 13]
)

# ── SECTION 7: NOTES ────────────────────────────────────────
add_section_header(doc, '7.  IMPORTANT NOTES TO PRINTER')
notes = [
    "⚠  Mockup images are AI-rendered references ONLY — recreate all artwork as vector files (AI or EPS) using exact color codes above.",
    "⚠  The Dharma Chakra wheel has intricate fine detail — minimum 300 DPI at final print size required.",
    "⚠  Gold areas should use metallic or shimmer finish ink if DTF printing allows.",
    "⚠  Request a physical strike-off sample (1 shirt) for client approval BEFORE bulk production.",
    "⚠  No bleed required — all designs contained within print area boundaries.",
]
for note in notes:
    p = doc.add_paragraph()
    run = p.add_run(note)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.italic = True

doc.add_paragraph()

# ── FOOTER ──────────────────────────────────────────────────
footer_table = doc.add_table(rows=1, cols=1)
footer_table.style = 'Table Grid'
cell = footer_table.cell(0, 0)
set_cell_bg(cell, '000000')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sachi Stream  —  Real People. Real Videos. No AI.  |  © 2026 LDNA Consulting LLC. All rights reserved.')
run.font.size = Pt(7)
run.font.color.rgb = GOLD

doc.save('Sachi_Stream_TShirt_PrintSpec.docx')
print("DOCX built successfully!")
